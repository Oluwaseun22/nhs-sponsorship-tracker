"""
cv_generator.py — CV Generation Module
"""
import logging
import os
import re
from datetime import datetime
from pathlib import Path
import config

log = logging.getLogger(__name__)

SKILL_KEYWORD_MAP = {
    "sql": "SQL querying & database management",
    "power bi": "Power BI reporting & dashboards",
    "python": "Python data analysis & scripting",
    "excel": "Advanced Microsoft Excel",
    "tableau": "Tableau data visualisation",
    "aws": "AWS cloud services",
    "azure": "Microsoft Azure",
    "data warehouse": "Data warehousing",
    "etl": "ETL pipeline development",
    "r ": "R statistical analysis",
    "machine learning": "Machine learning & predictive modelling",
    "nhs": "NHS data & information standards",
    "sus": "SUS & secondary uses data",
    "snowflake": "Snowflake cloud data platform",
    "databricks": "Databricks & Apache Spark",
    "dax": "DAX & data modelling",
    "reporting": "Operational reporting & analytics",
    "kpi": "KPI development & performance monitoring",
    "data quality": "Data quality & governance",
    "dashboard": "Dashboard design & visualisation",
}

def extract_keywords_from_jd(text: str) -> list:
    text_lower = text.lower()
    matched = []
    for keyword, label in SKILL_KEYWORD_MAP.items():
        if keyword in text_lower and label not in matched:
            matched.append(label)
        if len(matched) >= 3:
            break
    if len(matched) < 3:
        defaults = ["SQL querying & database management", "Power BI reporting & dashboards", "Advanced Microsoft Excel"]
        for d in defaults:
            if d not in matched:
                matched.append(d)
            if len(matched) >= 3:
                break
    return matched[:3]

def _replace_in_paragraph(para, replacements: dict):
    full_text = "".join(run.text for run in para.runs)
    for key, val in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, val)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = full_text
            else:
                para.add_run(full_text)
            break

def generate_cv(job: dict, mode: str = None) -> str:
    from docx import Document
    mode = mode or config.CV_MODE
    template_path = config.CV_TEMPLATE_PATH
    output_dir = Path(config.CV_OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)

    jd_text = job.get("full_description", "") + " " + job.get("summary", "")
    skills = extract_keywords_from_jd(jd_text)

    if mode == "ai":
        personal_statement, tailored_bullet = _generate_ai_content(job, skills)
    else:
        personal_statement = (
            f"A results-driven data professional with hands-on experience in SQL, Power BI, and Python, "
            f"seeking to contribute to {job.get('employer', 'the NHS')} as a {job.get('title', 'Data Analyst')}. "
            f"Experienced in delivering operational insights and data quality improvements within complex organisations. "
            f"Eligible for Certificate of Sponsorship under the Skilled Worker visa route."
        )
        tailored_bullet = (
            f"Delivered data analysis and reporting outputs aligned to the requirements of a "
            f"{job.get('title', 'data analyst')} role, using SQL, Power BI, and Excel to support decision-making."
        )

    title_clean = re.sub(r"\s*[\(\-]\s*(Band|AfC).*", "", job.get("title", "Data Analyst")).strip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Segun_Toriola_CV_{title_clean.replace(' ', '_')}_{timestamp}.docx"
    output_path = output_dir / filename

    replacements = {
        "{{JOB_TITLE}}": title_clean,
        "{{PERSONAL_STATEMENT}}": personal_statement,
        "{{SKILL_1}}": skills[0],
        "{{SKILL_2}}": skills[1],
        "{{SKILL_3}}": skills[2],
        "{{TAILORED_BULLET}}": tailored_bullet,
    }

    doc = Document(template_path)
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    doc.save(str(output_path))
    log.info(f"CV generated: {output_path}")
    return str(output_path)

def _generate_ai_content(job: dict, skills: list):
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
        prompt = (
            f"Write a 3-sentence personal statement for a CV applying to: {job.get('title')} at {job.get('employer')}.\n"
            f"Key skills to mention: {', '.join(skills)}.\n"
            f"End with: 'Eligible for Certificate of Sponsorship under the Skilled Worker visa route.'\n"
            f"Then write one bullet point (starting with a past-tense verb) for a work experience section.\n"
            f"Format: STATEMENT: <text>\nBULLET: <text>"
        )
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}]
        )
        response = message.content[0].text
        statement = re.search(r"STATEMENT:\s*(.+?)(?=BULLET:|$)", response, re.DOTALL)
        bullet = re.search(r"BULLET:\s*(.+)", response, re.DOTALL)
        return (
            statement.group(1).strip() if statement else "",
            bullet.group(1).strip() if bullet else "",
        )
    except Exception as e:
        log.warning(f"AI generation failed ({e}), using basic mode")
        return _generate_ai_content.__wrapped__(job, skills) if hasattr(_generate_ai_content, '__wrapped__') else ("", "")
